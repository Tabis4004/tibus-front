#!/usr/bin/env python3
"""Génère le manuel d'utilisation Tibus (DOCX + illustrations PNG)."""

from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUTPUT_DOCX = ROOT / "Manuel_Utilisation_Tibus_Compagnie.docx"

# Couleurs Tibus (approximation)
PRIMARY = (37, 99, 235)
BG = (248, 250, 252)
SIDEBAR = (15, 23, 42)
SIDEBAR_TEXT = (226, 232, 240)
MUTED = (100, 116, 139)
WHITE = (255, 255, 255)


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_card(draw, x, y, w, h, title, lines, active=False):
    fill = (239, 246, 255) if active else WHITE
    draw.rounded_rectangle((x, y, x + w, y + h), radius=12, fill=fill, outline=(203, 213, 225))
    draw.text((x + 14, y + 12), title, fill=PRIMARY if active else (15, 23, 42), font=load_font(13, True))
    ly = y + 36
    for line in lines:
        draw.text((x + 14, ly), line, fill=MUTED, font=load_font(11))
        ly += 18


def make_owner_sidebar_image(path: Path):
    w, h = 900, 620
    img = Image.new("RGB", (w, h), BG)
    draw = ImageDraw.Draw(img)

    draw.rectangle((0, 0, 240, h), fill=SIDEBAR)
    draw.text((20, 24), "TIBUS — Owner", fill=WHITE, font=load_font(14, True))
    draw.text((20, 48), "Ma Compagnie", fill=SIDEBAR_TEXT, font=load_font(11))

    sections = [
        ("GÉNÉRAL", ["▸ Aperçu", "  Mon entreprise", "  Avis", "  Codes promo", "  Abonnement"]),
        ("ANALYSE", ["  Analyses", "  Rapports Billets", "  Rapports Voyages", "  Voyageurs", "  Journal ventes", "  Caisse guichet"]),
        ("PARAMÉTRAGE", ["  Contact", "  Fidélité", "  Colis", "  Fond garantie", "  Pénalités"]),
        ("OPÉRATION", ["  Flotte", "  Gares", "  Itinéraires", "  Voyages", "  Scanner", "  Équipe"]),
    ]
    y = 90
    for title, items in sections:
        draw.text((20, y), title, fill=(148, 163, 184), font=load_font(9, True))
        y += 18
        for item in items:
            color = PRIMARY if item.startswith("▸") else SIDEBAR_TEXT
            draw.text((24, y), item.replace("▸ ", ""), fill=color, font=load_font(11, item.startswith("▸")))
            y += 22
        y += 6

    draw_card(draw, 270, 40, 600, 120, "Aperçu — Tableau de bord", [
        "Indicateurs clés : ventes, réservations, occupation",
        "Accès rapide aux actions du jour",
    ], active=True)
    draw_card(draw, 270, 180, 290, 100, "Voyages du jour", ["Départs programmés", "Places disponibles"])
    draw_card(draw, 580, 180, 290, 100, "Revenus", ["Guichet + en ligne", "Par période"])
    draw.text((270, 310), "Illustration — Espace Owner (desktop)", fill=MUTED, font=load_font(10))

    img.save(path)


def make_seller_image(path: Path):
    w, h = 900, 560
    img = Image.new("RGB", (w, h), BG)
    draw = ImageDraw.Draw(img)

    draw.rounded_rectangle((20, 20, w - 20, 70), radius=10, fill=WHITE, outline=(226, 232, 240))
    draw.text((36, 36), "Espace guichet — Vendeur", fill=(15, 23, 42), font=load_font(15, True))

    for i, (label, val) in enumerate([
        ("Mode caisse", "Guichet"),
        ("Départs", "12"),
        ("Places", "84"),
        ("Commissions", "15 000 XOF"),
    ]):
        x = 30 + i * 210
        draw_card(draw, x, 90, 190, 80, label, [val])

    draw_card(draw, 30, 190, 400, 150, "Session caisse journalière", [
        "① Ouvrir la caisse le matin (fond de roulement)",
        "② Vendre des billets cash",
        "③ Reversement comptable en fin de service",
    ], active=True)

    draw_card(draw, 450, 190, 420, 320, "Départs disponibles", [
        "Lomé → Kara · 08:00 · 32/45 places",
        "[ Vente guichet ]",
        "",
        "Formulaire : nom, téléphone, siège, colis",
        "Impression du ticket + QR code",
    ])

    draw.rounded_rectangle((20, h - 70, w - 20, h - 20), radius=10, fill=SIDEBAR)
    for i, label in enumerate(["Accueil", "Scanner", "Guichet"]):
        x = 120 + i * 280
        draw.text((x, h - 52), label, fill=WHITE if i == 2 else SIDEBAR_TEXT, font=load_font(11, i == 2))

    img.save(path)


def make_company_staff_image(path: Path):
    w, h = 900, 520
    img = Image.new("RGB", (w, h), BG)
    draw = ImageDraw.Draw(img)

    draw.text((30, 24), "Espace équipe compagnie", fill=(15, 23, 42), font=load_font(16, True))
    draw.text((30, 52), "Comptable · Contrôleur", fill=MUTED, font=load_font(12))

    draw_card(draw, 30, 90, 260, 280, "Journal des ventes", [
        "Toutes ventes guichet + en ligne",
        "Filtres par date / canal",
        "Export des données",
    ], active=True)
    draw_card(draw, 310, 90, 260, 280, "Rapports voyages", [
        "Occupation par départ",
        "Performance trajets",
        "Suivi remplissage",
    ])
    draw_card(draw, 590, 90, 280, 280, "Validation caisse", [
        "Reversements vendeurs",
        "Approuver / rejeter",
        "Clôture session guichet",
    ])

    draw.rounded_rectangle((20, h - 70, w - 20, h - 20), radius=10, fill=SIDEBAR)
    for i, label in enumerate(["Accueil", "Scanner", "Ventes", "Voyages", "Caisse"]):
        x = 50 + i * 165
        draw.text((x, h - 52), label, fill=PRIMARY if label == "Caisse" else SIDEBAR_TEXT, font=load_font(10, label == "Caisse"))

    img.save(path)


def make_traveler_image(path: Path):
    w, h = 900, 500
    img = Image.new("RGB", (w, h), BG)
    draw = ImageDraw.Draw(img)

    draw.text((30, 24), "Espace voyageur", fill=(15, 23, 42), font=load_font(16, True))

    draw.rounded_rectangle((30, 70, w - 30, 160), radius=12, fill=WHITE, outline=(226, 232, 240))
    draw.text((50, 90), "Rechercher un trajet", fill=MUTED, font=load_font(11))
    draw.text((50, 115), "Départ · Destination · Date", fill=(148, 163, 184), font=load_font(12))

    draw_card(draw, 30, 180, 400, 140, "Réservation en ligne", [
        "Choix du siège",
        "Paiement Mobile Money (FedaPay)",
        "Billet QR dans « Mes réservations »",
    ], active=True)
    draw_card(draw, 450, 180, 420, 140, "Fidélité & parrainage", [
        "Points compagnie + plateforme",
        "Codes promo",
        "Parrainer un proche",
    ])

    draw_card(draw, 30, 340, 840, 100, "Après paiement", [
        "Ticket avec référence TB-XXXXXXXX et QR code · Présentez-le à l'embarquement",
    ])

    img.save(path)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
    return h


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 65, 85)
    run.bold = bold
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10.5)


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 16):
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()


def build_docx():
    ASSETS.mkdir(parents=True, exist_ok=True)
    owner_img = ASSETS / "fig-owner.png"
    seller_img = ASSETS / "fig-vendeur.png"
    staff_img = ASSETS / "fig-equipe.png"
    traveler_img = ASSETS / "fig-voyageur.png"

    make_owner_sidebar_image(owner_img)
    make_seller_image(seller_img)
    make_company_staff_image(staff_img)
    make_traveler_image(traveler_img)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_heading("Manuel d'utilisation Tibus", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Guide compagnie de transport — 4 rôles principaux")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph()
    add_para(
        doc,
        "Ce document décrit les menus et les actions de chaque profil utilisateur sur la plateforme Tibus. "
        "Il s'adresse aux gérants, vendeurs guichet, comptables, contrôleurs et, en annexe, aux voyageurs clients de votre compagnie.",
    )

    add_heading(doc, "1. Les quatre rôles principaux", 1)
    add_bullets(doc, [
        "Owner (gérant) — Configuration complète de la compagnie, flotte, voyages, équipe et pilotage.",
        "Vendeur (guichet) — Vente de billets cash, caisse journalière, scan embarquement, colis.",
        "Équipe compagnie (comptable / contrôleur) — Suivi des ventes, rapports voyages, validation des reversements caisse.",
        "Voyageur — Recherche, réservation en ligne, paiement Mobile Money, billet QR.",
    ])

    add_heading(doc, "2. Owner — Gérant de la compagnie", 1)
    add_para(doc, "Accès : menu latéral sur /fr/owner (desktop) ou menu hamburger (mobile).")
    add_image(doc, owner_img, "Figure 1 — Navigation Owner et tableau de bord")

    owner_menus = [
        ("Aperçu", "Tableau de bord : indicateurs ventes, réservations, occupation. Point de départ chaque journée."),
        ("Mon entreprise", "Profil public : nom, logo, description, téléphone, e-mail, site web."),
        ("Avis", "Lire et répondre aux avis laissés par les voyageurs."),
        ("Codes promo", "Créer des codes de réduction (% ou montant fixe), dates de validité, limite d'usage."),
        ("Abonnement", "Formule Tibus active, renouvellement et statut."),
        ("Analyses", "Vue synthétique revenus, réservations et tendances."),
        ("Rapports Billets", "Détail des ventes par période, trajet et canal (guichet / en ligne)."),
        ("Rapports Voyages", "Occupation et performance de chaque départ programmé."),
        ("Voyageurs", "Clients récurrents et historique d'achats."),
        ("Journal des ventes", "Liste complète des transactions avec filtres et annulations."),
        ("Caisse guichet", "Suivi des sessions caisse vendeurs et reversements."),
        ("Contact", "Numéros WhatsApp et canaux visibles par les voyageurs."),
        ("Fidélité", "Programme points compagnie : règles, activation, suivi."),
        ("Colis autonomes", "Module colis sans billet : enregistrement et suivi."),
        ("Fond de garantie", "Solde garantie pour réservations en ligne."),
        ("Pénalités annulation", "Règles de remboursement et pénalités."),
        ("Flotte", "Bus : immatriculation, capacité, type."),
        ("Gares", "Points d'arrêt et villes desservies."),
        ("Itinéraires", "Liaisons gare à gare avec tarifs et kilométrage."),
        ("Voyages", "Programmation des départs (date, heure, bus, capacité)."),
        ("Scanner billets", "Contrôle QR à l'embarquement."),
        ("Équipe", "Inviter vendeurs, comptables, contrôleurs ; attribuer les rôles."),
    ]
    for name, desc in owner_menus:
        add_para(doc, name, bold=True)
        add_para(doc, desc)

    add_heading(doc, "3. Vendeur — Guichet", 1)
    add_para(doc, "Accès : /fr/seller · Barre mobile : Accueil · Scanner · Guichet")
    add_image(doc, seller_img, "Figure 2 — Espace vendeur : caisse et vente")

    seller_flow = [
        ("Ouvrir la caisse", "Chaque matin : « Session caisse journalière » avec fond de roulement (souvent 0). Obligatoire avant toute vente cash."),
        ("Choisir un départ", "Liste des voyages avec places disponibles. Cliquer « Vente guichet »."),
        ("Émettre le billet", "Saisir nom, téléphone (fidélité), choisir le siège, colis optionnel. Un ticket = une référence TB- unique + QR."),
        ("Vente multi-voyageurs", "Ajouter plusieurs voyageurs : chacun reçoit son propre ticket."),
        ("Reversement", "En fin de service : soumettre le montant au comptable. La caisse passe en attente de validation."),
        ("Scanner", "À l'embarquement : scanner le QR, marquer le passager à bord, détecter les doublons."),
        ("Colis", "Si activé : onglet Colis pour enregistrer des envois autonomes."),
        ("Ventes compagnie", "Onglet historique des ventes de la compagnie (lecture + annulation si autorisé)."),
    ]
    for name, desc in seller_flow:
        add_para(doc, name, bold=True)
        add_para(doc, desc)

    add_heading(doc, "4. Équipe compagnie — Comptable & Contrôleur", 1)
    add_para(doc, "Accès : /fr/company/* · Le comptable valide les caisses ; le contrôleur scanne les billets.")
    add_image(doc, staff_img, "Figure 3 — Espace comptable / contrôleur")

    staff_menus = [
        ("Journal des ventes (/company/sales)", "Toutes les ventes de la compagnie. Filtrer, contrôler, exporter."),
        ("Rapports voyages (/company/trip-reports)", "Taux de remplissage, départs, performance par ligne."),
        ("Validation caisse (/company/cash-register)", "Approuver les reversements vendeurs. Clôture la session guichet du vendeur."),
        ("Scanner (/verify/scan)", "Contrôle embarquement : authenticité billet, statut payé, anti-fraude."),
        ("Fond de garantie (/company/guarantee-fund)", "Consultation solde garantie (selon droits)."),
    ]
    for name, desc in staff_menus:
        add_para(doc, name, bold=True)
        add_para(doc, desc)

    add_para(doc, "Important — Le comptable n'ouvre pas de caisse guichet. Seuls les vendeurs ouvrent leur session journalière.", bold=True)

    add_heading(doc, "5. Voyageur — Client final", 1)
    add_para(doc, "Accès : site public /fr — recherche sans compte ; réservation avec connexion.")
    add_image(doc, traveler_img, "Figure 4 — Parcours voyageur")

    traveler_steps = [
        ("Rechercher", "Ville départ, arrivée, date. Comparer les compagnies et horaires."),
        ("Réserver", "Choisir un départ → nom, téléphone, siège, réseau Mobile Money."),
        ("Payer", "Redirection FedaPay. Le siège n'est confirmé qu'après paiement réussi."),
        ("Mes réservations", "Billet avec référence et QR code à présenter à l'embarquement."),
        ("Fidélité", "Points compagnie et plateforme, codes promo, parrainage."),
    ]
    for name, desc in traveler_steps:
        add_para(doc, name, bold=True)
        add_para(doc, desc)

    add_heading(doc, "6. Flux quotidien recommandé", 1)
    add_bullets(doc, [
        "Matin — Owner vérifie les départs du jour · Vendeur ouvre sa caisse.",
        "Journée — Ventes guichet + réservations en ligne · Contrôleur scanne à l'embarquement.",
        "Soir — Vendeur soumet reversement · Comptable valide · Owner consulte les rapports.",
    ])

    add_heading(doc, "7. Aide intégrée", 1)
    add_para(
        doc,
        "Bouton « Explorer les fonctionnalités » (header, sidebar ou menu utilisateur) : guide interactif "
        "rejouable à tout moment. Navigation clavier : flèches ← →, Entrée, Échap.",
    )

    add_heading(doc, "8. Dépannage rapide", 1)
    troubles = [
        ("Vente guichet impossible", "Vérifier que la caisse est ouverte. Message d'erreur affiché en détail dans l'application."),
        ("Réservation complète", "Plus de places sur le départ — choisir un autre horaire."),
        ("Reversement bloqué", "Un reversement est déjà en attente ; attendre validation comptable."),
        ("Scanner invalide", "Billet annulé, non payé ou déjà embarqué — message explicite à l'écran."),
    ]
    for issue, fix in troubles:
        add_para(doc, f"{issue} — {fix}")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("\n— Document généré pour Tibus · Compagnie de transport —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(OUTPUT_DOCX)
    print(f"OK: {OUTPUT_DOCX}")
    print(f"Assets: {ASSETS}")


if __name__ == "__main__":
    build_docx()
