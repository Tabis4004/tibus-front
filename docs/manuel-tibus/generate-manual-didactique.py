#!/usr/bin/env python3
"""Manuel didactique Tibus — formation Owner (gérant compagnie)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
CAPTURES = ROOT / "captures"
OUT = ROOT / "Manuel_Didactique_Tibus_Owner.docx"

DARK = RGBColor(15, 23, 42)
BLUE = RGBColor(37, 99, 235)
GREY = RGBColor(100, 116, 139)
BODY = RGBColor(51, 65, 85)


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = DARK
    return p


def add_p(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10.5)
            r.font.color.rgb = BODY


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for r in p.runs:
            r.font.size = Pt(10.5)
            r.font.color.rgb = BODY


def add_img(doc, name, caption, width=16.0):
    path = CAPTURES / name
    if not path.exists():
        add_p(doc, f"[Capture manquante : {name}]", italic=True)
        return
    doc.add_picture(str(path), width=Cm(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = GREY
    doc.add_paragraph()


def add_menu_block(doc, title, objectif, acces, procedure, capture=None, astuce=None):
    add_p(doc, title, bold=True)
    add_p(doc, f"Objectif pédagogique : {objectif}")
    add_p(doc, f"Accès : {acces}")
    if capture:
        add_img(doc, capture, f"Figure — {title} (capture réelle · Tibus Démo Transport)")
    add_p(doc, "Procédure :", bold=True)
    add_bullets(doc, procedure)
    if astuce:
        add_p(doc, f"Astuce formateur : {astuce}", italic=True)
    doc.add_paragraph()


OWNER_MENUS = [
    ("Aperçu", "Comprendre la santé de l'activité en un coup d'œil.", "/fr/owner",
     ["Consulter ventes du jour, revenus, occupation.", "Identifier les départs sous-remplis.", "Utiliser les raccourcis d'action."],
     "owner-real-overview.png", "Commencer chaque journée par cet écran."),
    ("Mon entreprise", "Paramétrer l'identité publique de la compagnie.", "/fr/owner/company",
     ["Renseigner nom, logo, description.", "Ajouter téléphone, e-mail, site.", "Enregistrer — visible sur le profil voyageur."],
     "owner-real-company.png"),
    ("Avis", "Gérer la réputation en ligne.", "/fr/owner/reviews",
     ["Lire les avis récents.", "Répondre aux commentaires.", "Identifier les points d'amélioration service."],
     "owner-real-reviews.png"),
    ("Codes promo", "Stimuler les ventes par promotions.", "/fr/owner/promo-codes",
     ["Créer un code (% ou montant fixe).", "Définir validité et limite d'usage.", "Optionnel : lier à un trajet précis."],
     "owner-real-promo-codes.png"),
    ("Abonnement", "Piloter la formule Tibus.", "/fr/owner/subscription",
     ["Vérifier plan actif et date de fin.", "Comparer les offres.", "Renouveler si nécessaire."],
     "owner-real-subscription.png"),
    ("Analyses", "Vue synthétique performance.", "/fr/owner/analytics",
     ["Analyser tendances revenus et réservations.", "Comparer périodes.", "Décider ajustements capacité/prix."],
     "owner-real-analytics.png"),
    ("Rapports Billets", "Audit détaillé des ventes.", "/fr/owner/analytics/tickets",
     ["Filtrer par date, trajet, canal.", "Distinguer guichet vs en ligne.", "Exporter pour comptabilité."],
     "owner-real-analytics-tickets.png"),
    ("Rapports Voyages", "Suivre l'occupation par départ.", "/fr/owner/analytics/trips",
     ["Voir taux de remplissage.", "Repérer trajets rentables.", "Ajuster fréquence des départs."],
     "owner-real-analytics-trips.png"),
    ("Voyageurs", "Connaître la clientèle.", "/fr/owner/analytics/travelers",
     ["Clients récurrents.", "Historique achats.", "Base pour fidélité ciblée."],
     "owner-real-analytics-travelers.png"),
    ("Journal des ventes", "Traçabilité complète transactions.", "/fr/owner/sales",
     ["Liste toutes ventes.", "Annulation billet si autorisé.", "Réconciliation comptable."],
     "owner-real-sales.png"),
    ("Caisse guichet", "Superviser sessions vendeurs.", "/fr/owner/cash-register",
     ["Voir caisses ouvertes.", "Suivre reversements.", "Contrôler écarts de caisse."],
     "owner-real-cash-register.png"),
    ("Contact", "Canaux visibles voyageurs.", "/fr/owner/messages",
     ["Configurer WhatsApp.", "Messages d'accueil.", "Support client."],
     "owner-real-messages.png"),
    ("Fidélité", "Programme points compagnie.", "/fr/owner/loyalty",
     ["Activer le programme.", "Définir règles gain/usage.", "Suivre adoption."],
     "owner-real-loyalty.png"),
    ("Colis autonomes", "Module colis sans billet.", "/fr/owner/colis",
     ["Activer le module.", "Paramétrer tarifs.", "Former l'équipe guichet après activation."],
     "owner-real-colis.png"),
    ("Fond de garantie", "Sécuriser réservations en ligne.", "/fr/owner/guarantee-fund",
     ["Consulter solde.", "Comprendre retenues.", "Recharger si seuil bas."],
     "owner-real-guarantee-fund.png"),
    ("Pénalités annulation", "Règles remboursement.", "/fr/owner/cancellation-policy",
     ["Définir délais et % pénalité.", "Communiquer aux voyageurs.", "Appliqué automatiquement."],
     "owner-real-cancellation-policy.png"),
    ("Flotte", "Gérer les bus.", "/fr/owner/buses",
     ["Ajouter bus : modèle, immatriculation, capacité.", "Mettre à jour ou retirer.", "Capacité = sièges vendables."],
     "owner-real-buses.png"),
    ("Gares", "Points d'arrêt.", "/fr/owner/stations",
     ["Créer gares par ville.", "Nommer clairement (Gare — Ville).", "Utilisées dans itinéraires."],
     "owner-real-stations.png"),
    ("Itinéraires", "Liaisons et tarifs.", "/fr/owner/routes",
     ["Lier gare départ → arrivée.", "Fixer prix et kilométrage.", "Base de la programmation."],
     "owner-real-routes.png"),
    ("Voyages", "Programmer les départs.", "/fr/owner/trips",
     ["Choisir itinéraire et bus.", "Date, heure, capacité.", "Publie le départ à la vente."],
     "owner-real-trips.png"),
    ("Scanner billets", "Contrôle embarquement (accessible depuis le menu Owner).", "/fr/verify/scan",
     ["Scanner QR passager.", "Vérifier authenticité.", "Marquer embarqué — anti-doublon."],
     "owner-real-scan.png"),
    ("Équipe", "Gérer les accès.", "/fr/owner/sellers",
     ["Inviter par e-mail.", "Attribuer rôle : vendeur, comptable, contrôleur.", "Révoquer si départ."],
     "owner-real-sellers.png"),
]


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    t = doc.add_heading("Manuel didactique Tibus", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Formation Owner — gérant de compagnie")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = BLUE
    doc.add_paragraph()
    add_p(doc, "Document de formation pour le gérant (owner) : menus, paramétrage, pilotage et supervision.", italic=True)
    add_p(doc, "Captures : compte tabiscompany@gmail.com · compagnie Tibus Démo Transport.", italic=True)
    doc.add_page_break()

    add_h(doc, "Sommaire", 1)
    for line in [
        "1. Introduction — Rôle Owner dans Tibus",
        "2. Connexion et navigation",
        "3. Les 22 menus Owner (fiches détaillées)",
        "4. Mise en route d'une compagnie",
        "5. Pilotage au quotidien",
        "6. Cycle de vie d'un billet (vue Owner)",
        "7. Aide intégrée et dépannage",
        "8. Glossaire Owner",
    ]:
        add_p(doc, line)
    doc.add_page_break()

    add_h(doc, "1. Introduction — Rôle Owner dans Tibus", 1)
    add_p(doc,
        "L'Owner est le gérant de la compagnie sur Tibus. Il configure l'offre (gares, itinéraires, bus, voyages), "
        "pilote la performance (analyses, ventes, caisse), règle les politiques commerciales (promos, annulations, fidélité) "
        "et invite son équipe (vendeurs, comptables, contrôleurs).")
    add_p(doc, "Ce manuel couvre uniquement l'espace Owner (/fr/owner). Les autres rôles sont créés et supervisés depuis le menu Équipe.", bold=True)
    add_p(doc, "Objectifs pédagogiques :", bold=True)
    add_bullets(doc, [
        "Maîtriser les 22 menus du tableau de bord Owner.",
        "Enchaîner la mise en route : gares → itinéraires → flotte → voyages.",
        "Superviser ventes, caisse et rapports sans quitter l'interface Owner.",
        "Former une équipe autonome en configurant correctement les accès.",
    ])
    add_img(doc, "owner-real-overview.png", "Figure 1.1 — Aperçu Owner (capture réelle)")

    add_h(doc, "2. Connexion et navigation", 1)
    add_p(doc, "Accès Owner :", bold=True)
    add_numbered(doc, [
        "Ouvrir l'application Tibus et se connecter avec le compte gérant.",
        "Après connexion, accéder à /fr/owner (menu « Espace compagnie »).",
        "Le menu latéral gauche liste les 22 sections — chacune est détaillée au chapitre 3.",
        "Le guide interactif (« Explorer les fonctionnalités ») peut être relancé à tout moment.",
    ])
    add_p(doc, "Compte de démonstration utilisé pour les captures de ce manuel : tabiscompany@gmail.com · Tibus Démo Transport.", italic=True)

    add_h(doc, "3. Les 22 menus Owner", 1)
    add_p(doc, "Chaque menu dispose d'une fiche : objectif, URL, capture réelle et procédure pas à pas.", bold=True)
    for row in OWNER_MENUS:
        title, obj, acces, proc = row[0], row[1], row[2], row[3]
        capture = row[4] if len(row) > 4 else None
        astuce = row[5] if len(row) > 5 else None
        add_menu_block(doc, title, obj, acces, proc, capture=capture, astuce=astuce)
    doc.add_page_break()

    add_h(doc, "4. Mise en route d'une compagnie", 1)
    add_p(doc, "Ordre recommandé pour une nouvelle compagnie :", bold=True)
    add_numbered(doc, [
        "Mon entreprise : identité, contacts, logo.",
        "Gares : créer les villes desservies.",
        "Itinéraires : lier gares, fixer tarifs.",
        "Flotte : enregistrer les bus et capacités.",
        "Voyages : programmer les premiers départs.",
        "Équipe : inviter vendeurs et comptables.",
        "Contact, fidélité, codes promo : selon stratégie commerciale.",
    ])
    add_p(doc, "Sans gares, itinéraires, bus et voyages, aucune vente n'est possible — ni en ligne ni au guichet.", italic=True)

    add_h(doc, "5. Pilotage au quotidien", 1)
    add_p(doc, "Routine matinale Owner :", bold=True)
    add_bullets(doc, [
        "Aperçu : ventes du jour et départs à venir.",
        "Voyages : vérifier capacités et bus assignés.",
        "Caisse guichet : s'assurer que les vendeurs ont ouvert leur session.",
    ])
    add_p(doc, "Routine soirée :", bold=True)
    add_bullets(doc, [
        "Journal des ventes et rapports billets.",
        "Caisse guichet : contrôler reversements en attente.",
        "Analyses : tendances hebdomadaires.",
    ])

    add_h(doc, "6. Cycle de vie d'un billet (vue Owner)", 1)
    add_numbered(doc, [
        "Programmation : Owner crée le voyage (itinéraire + bus + date).",
        "Vente : guichet ou réservation en ligne — visible dans Journal des ventes.",
        "Émission : billet TB- avec QR après paiement confirmé.",
        "Embarquement : contrôle via Scanner billets (menu Owner).",
        "Annulation : pénalités selon Pénalités annulation (menu Owner).",
    ])

    add_h(doc, "7. Aide intégrée et dépannage", 1)
    add_bullets(doc, [
        "Guide interactif : bouton « Explorer les fonctionnalités » sur l'Aperçu Owner.",
        "Navigation clavier du guide : flèches, Entrée, Échap.",
        "Messages d'erreur : lire le détail affiché (ex. caisse non ouverte côté vendeur).",
        "Abonnement : vérifier que le plan compagnie est actif si des fonctions sont bloquées.",
    ])

    add_h(doc, "8. Glossaire Owner", 1)
    gloss = [
        ("Départ / Réservation", "Instance d'un trajet à une date/heure — unité vendable."),
        ("Itinéraire", "Liaison entre deux gares avec tarif de base."),
        ("Session caisse", "Période guichet d'un vendeur ; supervisée depuis Caisse guichet."),
        ("Reversement", "Remise espèces vendeur → comptable ; suivi Owner dans Caisse guichet."),
        ("Fond de garantie", "Réserve pour sécuriser les réservations en ligne."),
        ("verifyToken", "Jeton QR sécurisé du billet, vérifié au scan."),
    ]
    for term, defn in gloss:
        add_p(doc, term, bold=True)
        add_p(doc, defn)

    add_p(doc, "\n— Fin du manuel didactique Tibus Owner —", italic=True)
    doc.save(OUT)
    print(f"Généré : {OUT}")


if __name__ == "__main__":
    build()
