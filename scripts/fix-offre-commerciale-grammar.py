#!/usr/bin/env python3
"""Corrections orthographiques (forme uniquement) du modèle d'offre commerciale."""
from pathlib import Path

from docx import Document

SRC = Path(__file__).resolve().parent.parent / "docs" / "offre-commerciale-tibus-modele.docx"

REPLACEMENTS = [
    ("facilité la gestion", "faciliter la gestion"),
    ("accroitre", "accroître"),
    ("module special est crée", "module spécial est créé"),
    ("au fur à mesure qu'il vende", "au fur et à mesure qu'il vend"),
    ("au besoin devenir négatif", "peut, au besoin, devenir négatif"),
    ("Dans cet ordre d'idée", "Dans cet ordre d'idées"),
    ("approche modulaire:", "approche modulaire :"),
    ("independent", "indépendants"),
    ("Control Anti-Fraud", "Contrôle anti-fraude"),
    ("Control anti-fraud", "Contrôle anti-fraude"),
    ("annulations.Prérequis", "annulations. Prérequis"),
    ("Billeterrie", "Billetterie"),
    ("Frais de service reservation", "Frais de service réservation"),
    ("D Courrier", "D — Courrier"),
    ("role d'utilisateur", "rôle d'utilisateur"),
]


def fix_text(text: str) -> str:
    out = text
    for old, new in REPLACEMENTS:
        out = out.replace(old, new)
    return out


def main() -> None:
    doc = Document(SRC)
    for para in doc.paragraphs:
        if para.text:
            fixed = fix_text(para.text)
            if fixed != para.text:
                para.text = fixed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        fixed = fix_text(para.text)
                        if fixed != para.text:
                            para.text = fixed
    doc.save(SRC)
    print(f"Updated: {SRC}")


if __name__ == "__main__":
    main()
