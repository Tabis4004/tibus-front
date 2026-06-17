#!/usr/bin/env python3
"""Génère le modèle Word offre commerciale Tibus (3 pages)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "offre-commerciale-tibus-modele.docx"
BLANK = "………………………"


def shade_header_row(table, row_idx: int = 0) -> None:
    for cell in table.rows[row_idx].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8F0FE")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)


def field(doc: Document, label: str, width: int = 40) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label} : ").bold = True
    run = p.add_run("_" * width)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def blank(doc: Document) -> None:
    doc.add_paragraph()


def bullet(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # —— Page 1 : Lettre ——
    doc.add_heading("LETTRE COMMERCIALE", 0)
    blank(doc)
    field(doc, "À l'attention de", 35)
    field(doc, "Compagnie de transport", 35)
    field(doc, "Adresse", 50)
    blank(doc)
    field(doc, "Objet", 55)
    p = doc.add_paragraph(
        "Proposition de partenariat — Plateforme digitale Tibus pour la gestion des "
        "réservations, de la flotte et du pilotage d'activité"
    )
    blank(doc)
    field(doc, "Date", 20)
    field(doc, "Référence offre", 25)
    blank(doc)

    for text in [
        "Madame, Monsieur le Directeur,",
        "",
        "Nous avons l'honneur de vous adresser la présente proposition dans le cadre du "
        "déploiement de Tibus, solution cloud dédiée aux compagnies de transport routier "
        "en Afrique de l'Ouest.",
        "",
        "Face aux enjeux de digitalisation des ventes (guichets, paiement mobile, contrôle "
        "embarquement) et de visibilité sur l'activité (rapports, caisse, suivi des lignes), "
        "Tibus propose une approche modulaire : vous activez uniquement les briques "
        "correspondant à votre organisation, sans investissement lourd en infrastructure locale.",
        "",
        "Notre proposition s'articule autour de trois volets joints à ce document :",
    ]:
        if text:
            doc.add_paragraph(text)
        else:
            blank(doc)

    for item in [
        "une offre technique détaillant les modules et l'architecture ;",
        "une offre financière modulable selon vos choix fonctionnels ;",
        "les conditions de mise en service et d'accompagnement.",
    ]:
        bullet(doc, item)

    blank(doc)
    doc.add_paragraph(
        "Nous restons à votre entière disposition pour une démonstration personnalisée et "
        "l'ajustement des modules à votre réseau de gares, à votre volume de ventes et à "
        "vos équipes terrain."
    )
    blank(doc)
    doc.add_paragraph(
        "Dans l'attente de votre retour, nous vous prions d'agréer, Madame, Monsieur le "
        "Directeur, l'expression de nos salutations distinguées."
    )
    blank(doc)
    field(doc, "Pour Tibus / [Raison sociale émettrice]", 30)
    field(doc, "Nom et qualité du signataire", 30)
    field(doc, "Contact (tél. / e-mail)", 40)

    doc.add_page_break()

    # —— Page 2 : Technique ——
    doc.add_heading("OFFRE TECHNIQUE", 0)
    p = doc.add_paragraph()
    p.add_run("Plateforme Tibus — Architecture et modules fonctionnels").bold = True
    blank(doc)
    field(doc, "Prospect", 40)
    field(doc, "Date de validité de l'offre", 25)
    blank(doc)

    doc.add_heading("1. Contexte et périmètre", 2)
    doc.add_paragraph(
        "Tibus est une plateforme Software-as-a-Service (SaaS) hébergée dans le cloud. "
        "Chaque compagnie cliente dispose d'un espace isolé (données, utilisateurs, paramètres) "
        "accessible via navigateur web sur ordinateur, tablette ou smartphone."
    )
    blank(doc)

    doc.add_heading("2. Architecture technique (incluse dans l'abonnement)", 2)
    arch = doc.add_table(rows=5, cols=2)
    arch.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Couche", "Description"),
            ("Interface utilisateur", "Application web responsive — hébergement cloud (Vercel)"),
            ("Données & sécurité", "PostgreSQL, authentification, isolation par compagnie (RLS)"),
            ("Services métier", "Fonctions cloud : paiements, notifications, webhooks, contrôle billets"),
            ("Paiements (TabisPay)", "Passerelle Mobile Money — commissions selon réseau opérateur"),
        ]
    ):
        arch.rows[i].cells[0].text = a
        arch.rows[i].cells[1].text = b
    shade_header_row(arch)
    blank(doc)

    doc.add_heading("3. Modules proposés (sélection à cocher)", 2)
    modules = [
        ("A", "Réservations & ventes au guichet", "Vendeurs, sièges, billets QR, scan embarquement, paiement mobile."),
        ("B", "Flotte, lignes & programmation", "Bus, chauffeurs, gares, trajets, horaires. Prérequis du module A."),
        ("C", "Colis & courrier", "Vente colis au guichet, natures et tarifs colis."),
        ("D", "Portail voyageur (web)", "Recherche, réservation et paiement en ligne, espace voyageur."),
        ("E", "Pilotage, caisse & rapports", "Tableau de bord, rapports, caisse, dépenses, rôles direction."),
        ("F", "Options avancées (sur devis)", "Promo, fidélité, API partenaire, TPE, annulations."),
    ]
    mod = doc.add_table(rows=1 + len(modules), cols=4)
    mod.style = "Table Grid"
    for j, h in enumerate(["☐", "Module", "Intitulé", "Contenu fonctionnel"]):
        mod.rows[0].cells[j].text = h
    for i, (code, title, desc) in enumerate(modules, 1):
        mod.rows[i].cells[0].text = "☐"
        mod.rows[i].cells[1].text = code
        mod.rows[i].cells[2].text = title
        mod.rows[i].cells[3].text = desc
    shade_header_row(mod)
    blank(doc)

    doc.add_heading("4. Mise en service & accompagnement", 2)
    for item in [
        "Création de l'espace compagnie et paramétrage des rôles",
        "Import gares, lignes, bus et horaires (données client)",
        "Formation guichetiers : ___ session(s) de ___ heure(s)",
        "Formation direction : ___ session(s)",
        "Délai indicatif de mise en production : ___ semaines",
        "Support : ___ (canal, horaires, délai incidents bloquants)",
    ]:
        bullet(doc, item, 10)

    doc.add_heading("5. Prérequis côté compagnie", 2)
    for item in [
        "Connexion Internet stable sur les guichets",
        "Fichier gares / lignes / horaires / flotte",
        "Convention Mobile Money si paiement en ligne",
        "Référent technique et référent métier désignés",
    ]:
        bullet(doc, item, 10)

    doc.add_heading("6. Hors périmètre standard", 2)
    doc.add_paragraph(
        "App native iOS/Android marque blanche, ERP tiers, hébergement on-premise — sur devis séparé."
    )

    doc.add_page_break()

    # —— Page 3 : Financière ——
    doc.add_heading("OFFRE FINANCIÈRE", 0)
    p = doc.add_paragraph()
    p.add_run("Grille modulaire — montants en francs CFA (F CFA)").bold = True
    blank(doc)
    field(doc, "Prospect", 40)
    field(doc, "Durée d'engagement proposée", 20)
    field(doc, "Date de validité", 25)
    blank(doc)

    doc.add_heading("1. Grille tarifaire par module", 2)
    fin_rows = [
        ("A", "Réservations & ventes au guichet"),
        ("B", "Flotte, lignes & programmation"),
        ("C", "Colis & courrier"),
        ("D", "Portail voyageur (web)"),
        ("E", "Pilotage, caisse & rapports"),
        ("F", "Options avancées"),
        ("—", "Envoi colis avec notification SMS"),
        ("—", "TOTAL modules sélectionnés"),
    ]
    fin = doc.add_table(rows=1 + len(fin_rows), cols=4)
    fin.style = "Table Grid"
    for j, h in enumerate(["Module", "Désignation", "Mise en service (F CFA)", "Abonnement mensuel (F CFA)"]):
        fin.rows[0].cells[j].text = h
    for i, (code, label) in enumerate(fin_rows, 1):
        fin.rows[i].cells[0].text = code
        fin.rows[i].cells[1].text = label
        fin.rows[i].cells[2].text = BLANK
        fin.rows[i].cells[3].text = BLANK
    shade_header_row(fin)
    blank(doc)

    doc.add_heading("2. Pack complet (modules A + B + C + D + E)", 2)
    pack = doc.add_table(rows=5, cols=2)
    pack.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Élément", "Montant (F CFA)"),
            ("Mise en service pack complet", BLANK),
            ("Abonnement mensuel pack (avant remise)", BLANK),
            ("Remise pack complet (___ %)", BLANK),
            ("Abonnement mensuel net", BLANK),
        ]
    ):
        pack.rows[i].cells[0].text = a
        pack.rows[i].cells[1].text = b
    shade_header_row(pack)
    blank(doc)

    doc.add_heading("3. Conditions de facturation", 2)
    for item in [
        "Mise en service : ___ % à la commande, ___ % à la mise en production",
        "Abonnement : facturation mensuelle d'avance, échéance le ___ du mois",
        "Frais passerelle Mobile Money : à la charge de ___ (compagnie / voyageur / partagé)",
        "Taux indicatif commissions opérateur : ___ %",
        "Révision tarifaire : après ___ mois, préavis ___ jours",
    ]:
        bullet(doc, item, 10)

    blank(doc)
    doc.add_heading("4. Synthèse de l'offre retenue", 2)
    syn = doc.add_table(rows=6, cols=2)
    syn.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Poste", "Valeur"),
            ("Modules retenus", "A ☐  B ☐  C ☐  D ☐  E ☐  F ☐"),
            ("Total mise en service", f"{BLANK} F CFA"),
            ("Abonnement mensuel", f"{BLANK} F CFA / mois"),
            ("Durée du contrat", f"{BLANK} mois"),
            ("Date de démarrage souhaitée", BLANK),
        ]
    ):
        syn.rows[i].cells[0].text = a
        syn.rows[i].cells[1].text = b
    shade_header_row(syn)

    blank(doc)
    p = doc.add_paragraph()
    p.add_run("Bon pour accord").bold = True
    blank(doc)
    field(doc, "Nom et qualité (client)", 35)
    field(doc, "Date et signature", 35)
    field(doc, "Cachet de la compagnie", 35)

    blank(doc)
    note = doc.add_paragraph()
    run = note.add_run(
        "Document modèle Tibus — brouillon. Montants et durées à compléter avant envoi. "
        "Sans valeur contractuelle tant que non signé."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
